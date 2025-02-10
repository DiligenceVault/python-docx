from docx import Document

document = Document("sample.docx")

para = document.add_paragraph()
comment = para.mark_comment_start("testing comment", "Test Author", "TA", "2025-01-01 00:00:00")
para.add_run("Some random text with a comment. ")
para.mark_comment_end(comment.id)
para.add_run("This text doesn't have a comment")

document.save("add_comments.docx")


# p = document.add_paragraph("Hello world")
# parent_comment = p.add_comment(
#     "Testing Comment",
#     author="Test Author",
#     initials="TA",
#     date="2025-01-01 00:00:00",
#     resolved=True,
# )
# p.add_comment(
#     "Testing Comment Reply 1",
#     author="Test Author",
#     initials="TA",
#     date="2025-01-01 00:00:00",
#     resolved=True,
#     parent=parent_comment,
# )
# p.add_comment(
#     "Testing Comment 2",
#     author="Test Author",
#     initials="TA",
#     date="2025-01-01 00:00:00",
#     resolved=True,
#     parent=parent_comment,
# )

# p2 = document.add_paragraph()
# comment = p2.mark_comment_start("testing comment", "Test Author", "TA", "2025-01-01 00:00:00")
# c2 = p2.mark_comment_start(
#     "testing comment reply", "Test Author", "TA", "2025-01-01 00:00:00", parent=comment
# )

# p2.add_run("Some random text")

# p3 = document.add_paragraph("Another para")
# p3.mark_comment_end(comment.id)
# p3.mark_comment_end(c2.id)
# p3.add_run("Lorem ipsum dolor sit amet")



# print(document.element.xpath(f"//w:commentRangeEnd[@w:id='{comment.id}']"), f"//w:commentRangeEnd[@w:id='{comment.id}']", "test")
# print(comment.id, c2.id)
